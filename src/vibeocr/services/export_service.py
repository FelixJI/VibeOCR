"""导出服务

支持将 OCR 结果导出为 Markdown、HTML、Word、Excel、纯文本格式。
"""

import base64
import io
import logging
import re
from pathlib import Path

from vibeocr.models.ocr_result import OCRResult
from vibeocr.utils.markdown_converter import HTML_STYLE

logger = logging.getLogger(__name__)

DISCARDED_BLOCK_TYPES = frozenset({
    "header", "footer", "page_number", "page_footnote", "aside_text",
})


class ExportService:
    """OCR 结果导出服务"""

    SUPPORTED_FORMATS = ["markdown", "html", "txt", "docx", "xlsx"]

    @staticmethod
    def export(
        result: OCRResult,
        output_path: Path,
        fmt: str,
    ) -> bool:
        """导出单个结果到文件

        Args:
            result: OCR 结果
            output_path: 输出文件路径（含文件名）
            fmt: 导出格式 (markdown, html, txt, docx, xlsx)

        Returns:
            是否成功
        """
        exporters = {
            "markdown": ExportService._export_markdown,
            "html": ExportService._export_html,
            "txt": ExportService._export_txt,
            "docx": ExportService._export_docx,
            "xlsx": ExportService._export_xlsx,
        }

        exporter = exporters.get(fmt)
        if not exporter:
            logger.error("不支持的导出格式: %s", fmt)
            return False

        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            return exporter(result, output_path)
        except Exception as e:
            logger.error("导出失败 [%s -> %s]: %s", fmt, output_path, e)
            return False

    @staticmethod
    def get_output_filename(source_name: str, fmt: str) -> str:
        """根据源文件名和格式生成输出文件名"""
        stem = Path(source_name).stem
        ext_map = {
            "markdown": ".md",
            "html": ".html",
            "txt": ".txt",
            "docx": ".docx",
            "xlsx": ".xlsx",
        }
        return stem + ext_map.get(fmt, ".txt")

    @staticmethod
    def _export_markdown(result: OCRResult, output_path: Path) -> bool:
        """导出为 Markdown"""
        content = result.markdown_text or result.raw_text
        output_path.write_text(content, encoding="utf-8")

        # 保存图片到 images 子目录
        if result.images:
            img_dir = output_path.parent / (output_path.stem + "_images")
            img_dir.mkdir(parents=True, exist_ok=True)
            for name, data in result.images.items():
                if isinstance(data, bytes):
                    (img_dir / name).write_bytes(data)

        logger.info("导出 Markdown: %s", output_path)
        return True

    @staticmethod
    def _export_html(result: OCRResult, output_path: Path) -> bool:
        """导出为 HTML（内嵌 base64 图片）"""
        html_body = result.html_text or result.raw_text

        # 将 markdown 中的图片引用替换为 base64 内嵌
        if result.images:
            for img_name, data in result.images.items():
                if isinstance(data, bytes):
                    b64 = base64.b64encode(data).decode("ascii")
                    ext = Path(img_name).suffix.lstrip(".") or "png"
                    mime = f"image/{ext}" if ext != "jpg" else "image/jpeg"
                    data_uri = f"data:{mime};base64,{b64}"
                    # 替换 markdown 图片语法
                    html_body = html_body.replace(f"({img_name})", f"({data_uri})")
                    html_body = html_body.replace(f'src="{img_name}"', f'src="{data_uri}"')

        full_html = (
            "<!DOCTYPE html>\n<html lang='zh-CN'>\n<head>\n"
            "<meta charset='utf-8'>\n"
            "<meta name='viewport' content='width=device-width, initial-scale=1'>\n"
            f"<title>{output_path.stem}</title>\n"
            f"{HTML_STYLE}\n"
            "</head>\n<body>\n{html_body}\n</body>\n</html>"
        )

        output_path.write_text(full_html, encoding="utf-8")
        logger.info("导出 HTML: %s", output_path)
        return True

    @staticmethod
    def _export_txt(result: OCRResult, output_path: Path) -> bool:
        """导出为纯文本"""
        content = result.raw_text or result.markdown_text
        output_path.write_text(content, encoding="utf-8")
        logger.info("导出纯文本: %s", output_path)
        return True

    @staticmethod
    def _export_docx(result: OCRResult, output_path: Path) -> bool:
        """导出为 Word 文档"""
        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document()
        content_list = getattr(result, "content_list", [])

        if content_list:
            for block in content_list:
                block_type = block.get("type", "text")
                text = block.get("text", "")

                if block_type in DISCARDED_BLOCK_TYPES:
                    continue
                elif block_type == "title":
                    level = min(block.get("level", 1), 6)
                    doc.add_heading(text, level=level)
                elif block_type == "text":
                    text_level = block.get("text_level")
                    if text_level and 1 <= text_level <= 6:
                        doc.add_heading(text, level=text_level)
                    elif text:
                        doc.add_paragraph(text)
                elif block_type == "table":
                    table_captions = block.get("table_caption") or []
                    if table_captions:
                        doc.add_paragraph(" ".join(table_captions), style="Caption")
                    html = block.get("table_body", "") or block.get("html", "")
                    if html:
                        ExportService._add_html_table_to_docx(doc, html)
                    table_footnotes = block.get("table_footnote") or []
                    for fn in table_footnotes:
                        if fn:
                            p = doc.add_paragraph(fn)
                            for run in p.runs:
                                run.font.size = Pt(9)
                elif block_type in ("image", "figure"):
                    img_path = block.get("img_path", "")
                    caption = block.get("image_caption") or block.get("chart_caption") or []
                    images = result.images or {}
                    img_added = False
                    if img_path and img_path in images:
                        data = images[img_path]
                        if isinstance(data, bytes):
                            try:
                                doc.add_picture(io.BytesIO(data), width=Inches(5))
                                img_added = True
                            except Exception:
                                pass
                    if not img_added:
                        label = " ".join(caption) if caption else text
                        if label:
                            doc.add_paragraph(f"[图片: {label}]")
                elif block_type in ("equation", "interline_equation", "inline_equation"):
                    if text:
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.font.name = "Consolas"
                        run.font.size = Pt(11)
                elif block_type == "list":
                    items = block.get("list_items", [])
                    for item in items:
                        doc.add_paragraph(item, style="List Bullet")
                elif block_type == "code":
                    body = block.get("code_body", "")
                    if body:
                        doc.add_paragraph(body, style="No Spacing")
        else:
            # 回退：使用纯文本
            text = result.raw_text or result.markdown_text
            for line in text.split("\n"):
                doc.add_paragraph(line)

        doc.save(str(output_path))
        logger.info("导出 Word: %s", output_path)
        return True

    @staticmethod
    def _add_html_table_to_docx(doc, html: str) -> None:
        """从 HTML 表格提取数据并添加到 docx"""
        rows_data: list[list[str]] = []

        # 提取行
        tr_pattern = re.compile(r"<tr[^>]*>(.*?)</tr>", re.DOTALL | re.IGNORECASE)
        td_pattern = re.compile(r"<t[dh][^>]*>(.*?)</t[dh]>", re.DOTALL | re.IGNORECASE)

        for tr_match in tr_pattern.finditer(html):
            cells = []
            for td_match in td_pattern.finditer(tr_match.group(1)):
                cell_text = re.sub(r"<[^>]+>", "", td_match.group(1)).strip()
                cells.append(cell_text)
            if cells:
                rows_data.append(cells)

        if not rows_data:
            return

        max_cols = max(len(row) for row in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=max_cols)
        table.style = "Table Grid"

        for i, row in enumerate(rows_data):
            for j, cell_text in enumerate(row):
                if j < max_cols:
                    table.rows[i].cells[j].text = cell_text

    @staticmethod
    def _export_xlsx(result: OCRResult, output_path: Path) -> bool:
        """导出为 Excel"""
        from openpyxl import Workbook

        wb = Workbook()
        ws_text = wb.active
        if ws_text is None:
            ws_text = wb.create_sheet("Sheet")
        content_list = getattr(result, "content_list", [])

        if content_list:
            table_count = 0
            has_text = False

            for block in content_list:
                block_type = block.get("type", "text")
                text = block.get("text", "")

                if block_type in DISCARDED_BLOCK_TYPES:
                    continue
                elif block_type == "table":
                    table_captions = block.get("table_caption") or []
                    if table_captions:
                        if not has_text:
                            has_text = True
                            ws_text.title = "文本汇总"
                        ws_text.append([f"[表格标题] {' '.join(table_captions)}"])
                    html = block.get("table_body", "") or block.get("html", "")
                    if html:
                        table_count += 1
                        rows_data = ExportService._parse_html_table(html)
                        if rows_data:
                            ws = wb.create_sheet(title=f"表格 {table_count}")
                            for row_idx, row in enumerate(rows_data):
                                for col_idx, cell_text in enumerate(row):
                                    ws.cell(
                                        row=row_idx + 1,
                                        column=col_idx + 1,
                                        value=cell_text,
                                    )

                elif block_type == "title" and text:
                    if not has_text:
                        has_text = True
                        ws_text.title = "文本汇总"
                    ws_text.append([f"[标题] {text}"])

                elif block_type == "text" and text:
                    if not has_text:
                        has_text = True
                        ws_text.title = "文本汇总"
                    text_level = block.get("text_level")
                    if text_level:
                        ws_text.append([f"{'#' * text_level} {text}"])
                    else:
                        ws_text.append([text])

                elif block_type in ("image", "figure"):
                    if not has_text:
                        has_text = True
                        ws_text.title = "文本汇总"
                    caption = block.get("image_caption") or block.get("chart_caption") or []
                    label = " ".join(caption) if caption else text
                    if label:
                        ws_text.append([f"[图片: {label}]"])

                elif block_type == "equation" and text:
                    if not has_text:
                        has_text = True
                        ws_text.title = "文本汇总"
                    ws_text.append([f"[公式] {text}"])

                elif block_type == "list":
                    items = block.get("list_items", [])
                    if items:
                        if not has_text:
                            has_text = True
                            ws_text.title = "文本汇总"
                        for item in items:
                            ws_text.append([f"• {item}"])

                elif block_type == "code":
                    body = block.get("code_body", "")
                    if body:
                        if not has_text:
                            has_text = True
                            ws_text.title = "文本汇总"
                        ws_text.append([f"[代码] {body}"])

            if not has_text and table_count > 0:
                if "Sheet" in wb.sheetnames:
                    del wb["Sheet"]
        else:
            ws_text.title = "文本"
            text = result.raw_text or result.markdown_text
            for line in text.split("\n"):
                ws_text.append([line])

        wb.save(str(output_path))
        logger.info("导出 Excel: %s", output_path)
        return True

    @staticmethod
    def _parse_html_table(html: str) -> list[list[str]]:
        """从 HTML 表格提取数据"""
        rows_data: list[list[str]] = []

        tr_pattern = re.compile(r"<tr[^>]*>(.*?)</tr>", re.DOTALL | re.IGNORECASE)
        td_pattern = re.compile(r"<t[dh][^>]*>(.*?)</t[dh]>", re.DOTALL | re.IGNORECASE)

        for tr_match in tr_pattern.finditer(html):
            cells = []
            for td_match in td_pattern.finditer(tr_match.group(1)):
                cell_text = re.sub(r"<[^>]+>", "", td_match.group(1)).strip()
                cells.append(cell_text)
            if cells:
                rows_data.append(cells)

        return rows_data
