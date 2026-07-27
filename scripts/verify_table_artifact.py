"""Verify offline table fixtures and the table semantics carried by release wheels."""

from __future__ import annotations

import argparse
import email
import hashlib
import json
import os
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any

_REQUIRED_RUNTIME_WHEELS = {
    "vibeocr-contracts-py",
    "vibeocr-client-py",
    "vibeocr-backend",
}
_SHA256_LENGTH = 64


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _read_json(path: Path) -> dict[str, Any]:
    value = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(value, dict):
        raise RuntimeError(f"JSON object expected: {path}")
    return value


def verify_fixture_manifest(fixture_root: Path) -> dict[str, Any]:
    """Validate synthetic provider fixtures and their content hashes."""

    root = fixture_root.resolve()
    manifest = _read_json(root / "manifest.json")
    if manifest.get("fixture_contract_version") != 1:
        raise RuntimeError("unsupported table fixture contract version")
    if manifest.get("synthetic") is not True:
        raise RuntimeError("table fixtures must be explicitly marked synthetic")
    records = manifest.get("fixtures")
    if not isinstance(records, list) or not records:
        raise RuntimeError("table fixture manifest has no fixture records")

    names: set[str] = set()
    providers: list[str] = []
    expected_topologies: list[tuple[int, int, tuple[tuple[Any, ...], ...]]] = []
    for record in records:
        if not isinstance(record, dict):
            raise RuntimeError("table fixture record must be an object")
        name = record.get("name")
        relative = record.get("file")
        provider = record.get("provider")
        provider_version = record.get("provider_version")
        source_schema = record.get("source_schema")
        expected_hash = record.get("sha256")
        if (
            not isinstance(name, str)
            or not name
            or not isinstance(relative, str)
            or not relative
            or not isinstance(provider, str)
            or not provider
            or not isinstance(provider_version, str)
            or not provider_version
            or not isinstance(source_schema, str)
            or not source_schema
        ):
            raise RuntimeError("table fixture record has missing provenance fields")
        if name in names:
            raise RuntimeError(f"duplicate table fixture name: {name}")
        names.add(name)
        if (
            not isinstance(expected_hash, str)
            or len(expected_hash) != _SHA256_LENGTH
            or any(char not in "0123456789abcdef" for char in expected_hash)
        ):
            raise RuntimeError(f"fixture {name} has invalid sha256")
        path = (root / relative).resolve()
        try:
            path.relative_to(root)
        except ValueError as error:
            raise RuntimeError(f"fixture {name} escapes manifest directory") from error
        if not path.is_file():
            raise RuntimeError(f"fixture {name} is missing: {relative}")
        actual_hash = _sha256(path)
        if actual_hash != expected_hash:
            raise RuntimeError(
                f"fixture {name} SHA256 mismatch: expected {expected_hash}, got {actual_hash}"
            )
        payload = _read_json(path)
        if payload.get("name") != name or payload.get("synthetic") is not True:
            raise RuntimeError(
                f"fixture {name} is not an explicitly synthetic matching payload"
            )
        source = payload.get("source")
        expected = payload.get("expected")
        if not isinstance(source, dict) or not isinstance(expected, dict):
            raise RuntimeError(
                f"fixture {name} must contain source and expected objects"
            )
        if not isinstance(source.get("provider_payload"), dict):
            raise RuntimeError(f"fixture {name} provider_payload must be an object")
        if (
            not isinstance(source.get("legacy_html"), str)
            or "<table" not in source["legacy_html"]
        ):
            raise RuntimeError(f"fixture {name} has no table source HTML")
        canonical = expected.get("canonical_table")
        if not isinstance(canonical, dict):
            raise RuntimeError(f"fixture {name} has no canonical table expectation")
        # Import only after fixture structure/hash checks so diagnostics remain useful
        # in a minimal build environment.
        from vibeocr.contracts.tables import TableModelV1

        table = TableModelV1.from_payload(canonical)
        expected_topologies.append(
            (
                table.row_count,
                table.column_count,
                tuple(
                    (
                        cell.row,
                        cell.column,
                        cell.rowspan,
                        cell.colspan,
                        cell.text,
                        cell.is_header,
                    )
                    for cell in table.cells
                ),
            )
        )
        providers.append(provider)

    first = expected_topologies[0]
    if any(topology != first for topology in expected_topologies[1:]):
        raise RuntimeError(
            "provider fixtures do not share one canonical table topology"
        )
    return {
        "fixture_count": len(records),
        "synthetic": True,
        "providers": providers,
    }


def _wheel_distribution(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        metadata_name = next(
            (
                name
                for name in archive.namelist()
                if name.endswith(".dist-info/METADATA")
            ),
            None,
        )
        if metadata_name is None:
            raise RuntimeError(f"wheel has no METADATA: {path.name}")
        metadata = email.message_from_bytes(archive.read(metadata_name))
    return str(metadata["Name"]).lower().replace("_", "-")


def _find_wheels(root: Path) -> dict[str, Path]:
    wheels: dict[str, Path] = {}
    for wheel in sorted(root.rglob("*.whl")):
        distribution = _wheel_distribution(wheel)
        if distribution in wheels:
            raise RuntimeError(f"duplicate wheel distribution: {distribution}")
        wheels[distribution] = wheel
    missing = _REQUIRED_RUNTIME_WHEELS - set(wheels)
    if missing:
        raise RuntimeError(f"table runtime wheel set incomplete: {sorted(missing)}")
    return wheels


def _verify_product_manifest(root: Path, wheels: dict[str, Path]) -> None:
    manifests = list(root.rglob("product-manifest.json"))
    if not manifests:
        return
    if len(manifests) != 1:
        raise RuntimeError("release artifact has multiple product manifests")
    manifest_path = manifests[0]
    manifest = _read_json(manifest_path)
    records = manifest.get("python_wheels")
    if not isinstance(records, list):
        raise RuntimeError("product manifest has no python_wheels")
    wheels_by_filename = {wheel.name: wheel for wheel in wheels.values()}
    seen_filenames: set[str] = set()
    for record in records:
        if not isinstance(record, dict):
            raise RuntimeError("product manifest wheel record is invalid")
        wheel_name = record.get("file")
        expected_hash = record.get("sha256")
        if not isinstance(wheel_name, str) or wheel_name in seen_filenames:
            raise RuntimeError("product manifest wheel file is missing or duplicated")
        seen_filenames.add(wheel_name)
        wheel = wheels_by_filename.get(wheel_name)
        if wheel is None:
            raise RuntimeError(
                f"product manifest references unavailable wheel: {wheel_name}"
            )
        actual_distribution = _wheel_distribution(wheel)
        declared_distribution = record.get("distribution")
        if declared_distribution is not None and (
            not isinstance(declared_distribution, str)
            or declared_distribution.lower().replace("_", "-") != actual_distribution
        ):
            raise RuntimeError(f"product manifest distribution mismatch: {wheel_name}")
        if _sha256(wheel) != expected_hash:
            raise RuntimeError(f"product manifest wheel hash mismatch: {wheel_name}")
    if seen_filenames != set(wheels_by_filename):
        missing = sorted(set(wheels_by_filename) - seen_filenames)
        raise RuntimeError(
            f"product manifest does not bind every packaged wheel: {missing}"
        )


def _provider_fixture_bundle(fixture_root: Path) -> list[dict[str, Any]]:
    manifest = _read_json(fixture_root / "manifest.json")
    bundle: list[dict[str, Any]] = []
    for record in manifest["fixtures"]:
        fixture = _read_json(fixture_root / record["file"])
        bundle.append(
            {
                "name": fixture["name"],
                "provider_payload": fixture["source"]["provider_payload"],
                "canonical_table": fixture["expected"]["canonical_table"],
            }
        )
    return bundle


def _run_wheel_semantics(
    wheels: dict[str, Path], output_dir: Path, fixture_root: Path
) -> None:
    """Install exact wheels offline, then execute provider and export smokes."""

    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "provider-fixtures.json").write_text(
        json.dumps(
            _provider_fixture_bundle(fixture_root),
            ensure_ascii=False,
            sort_keys=True,
        ),
        encoding="utf-8",
    )
    code = r"""
import json
import zipfile
from importlib import resources
from pathlib import Path
from types import SimpleNamespace

from vibeocr.contracts.tables import TableCellV1, TableModelV1
from vibeocr.models.ocr_result import OCRResult
from vibeocr.services.export_service import ExportService
from vibeocr.tables.blocks import canonicalize_table_block
from vibeocr.tables.html_adapter import table_model_from_html

output = Path.cwd()
schema = json.loads(
    resources.files("vibeocr.contracts")
    .joinpath("schemas/table-v1.schema.json")
    .read_text(encoding="utf-8")
)
assert schema["properties"]["schema_version"]["const"] == 1
assert schema["required"] == [
    "schema_version",
    "table_id",
    "row_count",
    "column_count",
    "coordinate_space",
    "cells",
    "provenance",
]

class FixturePipeline:
    def __init__(self, responses):
        self._responses = responses

    def predict(self, **_kwargs):
        return iter(self._responses)

class FixtureService:
    def __init__(self, responses):
        self._pipeline = FixturePipeline(responses)

    def get_or_create_pipeline(self, _name):
        return self._pipeline

def result_from_fixture(fixture):
    name = fixture["name"]
    payload = fixture["provider_payload"]
    if name.startswith("mineru-"):
        from vibeocr.services.mineru_service import MinerUService

        service = MinerUService.__new__(MinerUService)
        return service._build_ocr_result(payload, "fixture.pdf", data=None)
    if name == "paddleocr-table-recognition":
        from vibeocr.core.pipelines.pipeline_table import (
            TABLE_RECOGNITION_SPEC,
            TableRecognitionOptions,
        )

        return TABLE_RECOGNITION_SPEC.recognize(
            FixtureService([payload]), None, TableRecognitionOptions()
        )
    if name == "paddleocr-pp-structure":
        from vibeocr.core.pipelines.pipeline_pp_structure import (
            PP_STRUCTURE_V3_SPEC,
            PPStructureV3Options,
        )

        response = dict(payload)
        response["parsing_res_list"] = [
            SimpleNamespace(**block) for block in payload["parsing_res_list"]
        ]
        return PP_STRUCTURE_V3_SPEC.recognize(
            FixtureService([response]), None, PPStructureV3Options()
        )
    if name == "paddleocr-vl":
        from vibeocr.core.pipelines.pipeline_paddlocr_vl import (
            PADDLEOCR_VL_SPEC,
            PaddleOCRVLOptions,
        )

        return PADDLEOCR_VL_SPEC.recognize(
            FixtureService([payload]), None, PaddleOCRVLOptions()
        )
    raise AssertionError(f"unhandled provider fixture: {name}")

provider_fixtures = json.loads(
    (output / "provider-fixtures.json").read_text(encoding="utf-8")
)
providers_seen = set()
provider_result = None
for fixture in provider_fixtures:
    provider_result = result_from_fixture(fixture)
    block = next(
        value
        for value in provider_result.content_list
        if value.get("type") == "table" and isinstance(value.get("table"), dict)
    )
    assert block["table"] == fixture["canonical_table"], fixture["name"]
    providers_seen.add(
        "mineru" if fixture["name"].startswith("mineru-") else fixture["name"]
    )
assert providers_seen == {
    "mineru",
    "paddleocr-table-recognition",
    "paddleocr-pp-structure",
    "paddleocr-vl",
}

table = TableModelV1(
    table_id="release-table",
    row_count=2,
    column_count=3,
    cells=(
        TableCellV1("r0c0", 0, 0, text="vertical", rowspan=2),
        TableCellV1("r0c1", 0, 1, text="horizontal", colspan=2),
        TableCellV1("r1c1", 1, 1, text="left"),
        TableCellV1("r1c2", 1, 2, text="right"),
    ),
)
block = canonicalize_table_block({"type": "table", "table": table.to_payload()}, table_id=table.table_id, pipeline="release-verifier")
result = OCRResult(raw_text="", html_text=block["table_body"], content_list=[block])
for extension in ("html", "xlsx", "docx"):
    path = output / f"table.{extension}"
    assert ExportService.export(result, path, extension), extension
assert table_model_from_html((output / "table.html").read_text(encoding="utf-8"), table_id="release-table").merged_ranges() == ((0, 0, 1, 0), (0, 1, 0, 2))
from openpyxl import load_workbook
sheet = load_workbook(output / "table.xlsx")["表格 1"]
assert {str(value) for value in sheet.merged_cells.ranges} == {"A1:A2", "B1:C1"}
from docx import Document
document = Document(output / "table.docx")
assert len(document.tables) == 1
document_table = document.tables[0]
assert document_table.cell(0, 0)._tc is document_table.cell(1, 0)._tc
assert document_table.cell(0, 1)._tc is document_table.cell(0, 2)._tc
with zipfile.ZipFile(output / "table.docx") as archive:
    document_xml = archive.read("word/document.xml").decode("utf-8")
assert '<w:gridSpan w:val="2"' in document_xml
assert '<w:vMerge w:val="restart"' in document_xml
assert "<w:vMerge/>" in document_xml
"""
    with tempfile.TemporaryDirectory(
        prefix="vibeocr-table-wheel-install-"
    ) as temporary:
        install_root = Path(temporary)
        installed = subprocess.run(
            [
                sys.executable,
                "-m",
                "pip",
                "install",
                "--disable-pip-version-check",
                "--no-index",
                "--no-deps",
                "--target",
                str(install_root),
                *(str(path) for path in wheels.values()),
            ],
            text=True,
            capture_output=True,
            check=False,
        )
        (output_dir / "wheel-install.stdout.log").write_text(
            installed.stdout, encoding="utf-8"
        )
        (output_dir / "wheel-install.stderr.log").write_text(
            installed.stderr, encoding="utf-8"
        )
        if installed.returncode:
            raise RuntimeError(f"release wheel install failed; see {output_dir}")

        environment = os.environ.copy()
        environment["PYTHONPATH"] = str(install_root)
        environment["PYTHONNOUSERSITE"] = "1"
        completed = subprocess.run(
            [sys.executable, "-c", code],
            cwd=output_dir,
            env=environment,
            text=True,
            capture_output=True,
            check=False,
        )
        (output_dir / "wheel-semantic.stdout.log").write_text(
            completed.stdout, encoding="utf-8"
        )
        (output_dir / "wheel-semantic.stderr.log").write_text(
            completed.stderr, encoding="utf-8"
        )
        if completed.returncode:
            raise RuntimeError(
                f"release wheel table semantics failed; see {output_dir}"
            )


def verify_table_artifact(
    artifact: Path,
    report_dir: Path,
    fixture_root: Path = Path("tests/fixtures/table_contract/v1"),
) -> None:
    """Verify a wheel directory or packaged artifact before release upload."""

    target = artifact.resolve()
    if not target.exists():
        raise FileNotFoundError(f"table artifact target not found: {target}")
    with tempfile.TemporaryDirectory(prefix="vibeocr-table-artifact-") as temporary:
        root = Path(temporary)
        if target.is_dir():
            root = target
            wheels = _find_wheels(root)
            _verify_product_manifest(root, wheels)
            _run_wheel_semantics(wheels, report_dir, fixture_root)
            return
        if target.suffix.lower() == ".whl":
            wheels = _find_wheels(target.parent)
            _run_wheel_semantics(wheels, report_dir, fixture_root)
            return
        with zipfile.ZipFile(target) as archive:
            archive.extractall(root)
        wheels = _find_wheels(root)
        _verify_product_manifest(root, wheels)
        _run_wheel_semantics(wheels, report_dir, fixture_root)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("artifact", nargs="?", type=Path)
    parser.add_argument(
        "--fixture-root",
        type=Path,
        default=Path("tests/fixtures/table_contract/v1"),
    )
    parser.add_argument(
        "--report-dir", type=Path, default=Path("reports/table-artifact")
    )
    args = parser.parse_args()
    report = verify_fixture_manifest(args.fixture_root)
    print(json.dumps(report, ensure_ascii=False, sort_keys=True))
    if args.artifact is not None:
        verify_table_artifact(args.artifact, args.report_dir, args.fixture_root)
        print(f"table artifact verifier: PASS ({args.artifact})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
