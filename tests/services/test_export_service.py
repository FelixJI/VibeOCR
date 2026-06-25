"""导出服务测试"""

import io

import pytest

from vibeocr.models.ocr_result import OCRResult
from vibeocr.services.export_service import ExportService


def _make_ocr_result(
    content_list=None,
    images=None,
    raw_text="",
    markdown_text="",
    html_text="",
) -> OCRResult:
    return OCRResult(
        raw_text=raw_text,
        markdown_text=markdown_text,
        html_text=html_text,
        content_list=content_list or [],
        images=images or {},
    )


# --- 表格导出 ---


class TestTableExport:
    """表格导出修复：content_list 表格块使用 table_body 字段，而非 html"""

    def test_docx_table_from_table_body(self, tmp_path):
        content_list = [
            {"type": "text", "text": "标题文本"},
            {
                "type": "table",
                "table_body": "<tr><td>A</td><td>B</td></tr><tr><td>1</td><td>2</td></tr>",
            },
        ]
        result = _make_ocr_result(content_list=content_list)
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

        from docx import Document  # type: ignore[import-not-found]

        doc = Document(str(out))
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "A"
        assert doc.tables[0].rows[1].cells[1].text == "2"

    def test_xlsx_table_from_table_body(self, tmp_path):
        content_list = [
            {"type": "table", "table_body": "<tr><td>X</td><td>Y</td></tr>"},
        ]
        result = _make_ocr_result(content_list=content_list)
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")

        from openpyxl import load_workbook

        wb = load_workbook(str(out))
        assert "表格 1" in wb.sheetnames
        ws = wb["表格 1"]
        assert ws.cell(1, 1).value == "X"
        assert ws.cell(1, 2).value == "Y"

    def test_docx_table_fallback_to_html_field(self, tmp_path):
        """兼容：仍支持 html 字段（PaddleX 可能使用）"""
        content_list = [
            {"type": "table", "html": "<tr><td>H</td></tr>"},
        ]
        result = _make_ocr_result(content_list=content_list)
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

        from docx import Document  # type: ignore[import-not-found]

        doc = Document(str(out))
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "H"

    def test_docx_table_prefers_table_body_over_html(self, tmp_path):
        content_list = [
            {
                "type": "table",
                "table_body": "<tr><td>body</td></tr>",
                "html": "<tr><td>fallback</td></tr>",
            },
        ]
        result = _make_ocr_result(content_list=content_list)
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

        from docx import Document  # type: ignore[import-not-found]

        doc = Document(str(out))
        assert doc.tables[0].rows[0].cells[0].text == "body"


# --- 图片导出 ---


class TestImageExport:
    """图片导出修复：使用 img_path 字段匹配图片"""

    def test_docx_image_by_img_path(self, tmp_path):
        try:
            from PIL import Image as PILImage
        except ImportError:
            pytest.skip("PIL not installed")

        # 创建一个最小的有效 PNG
        img = PILImage.new("RGB", (10, 10), "red")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        img_bytes = buf.getvalue()

        content_list = [
            {"type": "image", "img_path": "images/img_0.png"},
        ]
        result = _make_ocr_result(
            content_list=content_list,
            images={"images/img_0.png": img_bytes},
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

        from docx import Document  # type: ignore[import-not-found]

        doc = Document(str(out))
        assert len(doc.inline_shapes) == 1

    def test_docx_image_fallback_to_caption(self, tmp_path):
        """图片无法嵌入时，回退到 caption 文本"""
        content_list = [
            {"type": "image", "img_path": "missing.png", "image_caption": ["Figure 1"]},
        ]
        result = _make_ocr_result(content_list=content_list, images={})
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

        from docx import Document  # type: ignore[import-not-found]

        doc = Document(str(out))
        assert any("Figure 1" in p.text for p in doc.paragraphs)

    def test_xlsx_image_uses_caption(self, tmp_path):
        content_list = [
            {"type": "image", "img_path": "img.png", "image_caption": ["Fig A"]},
        ]
        result = _make_ocr_result(content_list=content_list)
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")

        from openpyxl import load_workbook

        wb = load_workbook(str(out))
        ws = wb["文本汇总"]
        assert any(
            "Fig A" in str(cell.value or "") for row in ws.iter_rows() for cell in row
        )


# --- 标题/文本级别 ---


class TestTextLevelExport:
    """MinerU content_list 使用 text_level 字段表示标题级别"""

    def test_docx_text_level_heading(self, tmp_path):
        content_list = [
            {"type": "text", "text": "第一章", "text_level": 1},
            {"type": "text", "text": "普通段落"},
            {"type": "text", "text": "子标题", "text_level": 2},
        ]
        result = _make_ocr_result(content_list=content_list)
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

        from docx import Document  # type: ignore[import-not-found]

        doc = Document(str(out))
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) == 2
        assert headings[0].style.name == "Heading 1"
        assert headings[1].style.name == "Heading 2"

    def test_xlsx_text_level_prefix(self, tmp_path):
        content_list = [
            {"type": "text", "text": "标题一", "text_level": 1},
            {"type": "text", "text": "内容"},
        ]
        result = _make_ocr_result(content_list=content_list)
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")

        from openpyxl import load_workbook

        wb = load_workbook(str(out))
        ws = wb.active
        assert ws is not None
        values = [cell.value for row in ws.iter_rows() for cell in row if cell.value]
        assert "# 标题一" in values
        assert "内容" in values


# --- 其他块类型 ---


class TestOtherBlockTypes:
    """列表、公式、代码块导出"""

    def test_docx_list_items(self, tmp_path):
        content_list = [
            {"type": "list", "list_items": ["苹果", "香蕉"]},
        ]
        result = _make_ocr_result(content_list=content_list)
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

        from docx import Document  # type: ignore[import-not-found]

        doc = Document(str(out))
        list_items = [
            p.text for p in doc.paragraphs if "苹果" in p.text or "香蕉" in p.text
        ]
        assert len(list_items) == 2

    def test_xlsx_equation(self, tmp_path):
        content_list = [
            {"type": "equation", "text": "E=mc^2"},
        ]
        result = _make_ocr_result(content_list=content_list)
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")

        from openpyxl import load_workbook

        wb = load_workbook(str(out))
        ws2 = wb.active
        assert ws2 is not None
        values = [cell.value for row in ws2.iter_rows() for cell in row if cell.value]
        assert any("E=mc^2" in str(v) for v in values)


# --- 同名文件自动重命名 ---


class TestUniqueOutputPath:
    """get_unique_output_path 避免同名覆盖"""

    def test_no_conflict(self, tmp_path):
        path = tmp_path / "report.md"
        assert ExportService.get_unique_output_path(path) == path

    def test_one_conflict(self, tmp_path):
        (tmp_path / "report.md").write_text("old")
        result = ExportService.get_unique_output_path(tmp_path / "report.md")
        assert result == tmp_path / "report_1.md"

    def test_multiple_conflicts(self, tmp_path):
        (tmp_path / "report.md").write_text("a")
        (tmp_path / "report_1.md").write_text("b")
        (tmp_path / "report_2.md").write_text("c")
        result = ExportService.get_unique_output_path(tmp_path / "report.md")
        assert result == tmp_path / "report_3.md"

    def test_gap_in_sequence(self, tmp_path):
        (tmp_path / "report.md").write_text("a")
        (tmp_path / "report_2.md").write_text("b")
        result = ExportService.get_unique_output_path(tmp_path / "report.md")
        assert result == tmp_path / "report_1.md"
