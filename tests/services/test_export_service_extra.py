"""export_service 补充测试 — 覆盖 txt/html/markdown 导出、不支持的格式、get_output_filename 等"""

from vibeocr.contracts.tables import TableCellV1, TableModelV1
from vibeocr.models.ocr_result import OCRResult
from vibeocr.services.export_service import ExportService


def _make_result(**kwargs):
    return OCRResult(
        raw_text=kwargs.get("raw_text", ""),
        markdown_text=kwargs.get("markdown_text", ""),
        html_text=kwargs.get("html_text", ""),
        content_list=kwargs.get("content_list", []),
        images=kwargs.get("images", {}),
    )


def _mixed_table_payload() -> dict:
    return TableModelV1(
        table_id="canonical-table",
        row_count=2,
        column_count=3,
        cells=(
            TableCellV1(
                cell_id="a", row=0, column=0, rowspan=2, text="A"
            ),
            TableCellV1(
                cell_id="b", row=0, column=1, colspan=2, text="B"
            ),
            TableCellV1(cell_id="c", row=1, column=1, text="C"),
            TableCellV1(cell_id="d", row=1, column=2, text="D"),
        ),
    ).to_payload()


class TestExportUnsupportedFormat:
    def test_unknown_format_returns_false(self, tmp_path):
        result = _make_result(raw_text="hello")
        out = tmp_path / "test.xyz"
        assert not ExportService.export(result, out, "pdf")


class TestGetOutputFilename:
    def test_markdown(self):
        assert (
            ExportService.get_output_filename("report.pdf", "markdown") == "report.md"
        )

    def test_html(self):
        assert ExportService.get_output_filename("doc.txt", "html") == "doc.html"

    def test_txt(self):
        assert ExportService.get_output_filename("file.docx", "txt") == "file.txt"

    def test_docx(self):
        assert ExportService.get_output_filename("scan.png", "docx") == "scan.docx"

    def test_xlsx(self):
        assert ExportService.get_output_filename("data.pdf", "xlsx") == "data.xlsx"

    def test_unknown_format_falls_back_to_txt(self):
        assert ExportService.get_output_filename("file.pdf", "xyz") == "file.txt"


class TestExportTxt:
    def test_exports_raw_text(self, tmp_path):
        result = _make_result(raw_text="hello world")
        out = tmp_path / "test.txt"
        assert ExportService.export(result, out, "txt")
        assert out.read_text(encoding="utf-8") == "hello world"

    def test_exports_markdown_as_fallback(self, tmp_path):
        result = _make_result(markdown_text="# Title")
        out = tmp_path / "test.txt"
        assert ExportService.export(result, out, "txt")
        assert out.read_text(encoding="utf-8") == "# Title"


class TestExportMarkdown:
    def test_exports_markdown_text(self, tmp_path):
        result = _make_result(markdown_text="## Hello")
        out = tmp_path / "test.md"
        assert ExportService.export(result, out, "markdown")
        assert out.read_text(encoding="utf-8") == "## Hello"

    def test_falls_back_to_raw_text(self, tmp_path):
        result = _make_result(raw_text="plain text")
        out = tmp_path / "test.md"
        assert ExportService.export(result, out, "markdown")
        assert out.read_text(encoding="utf-8") == "plain text"

    def test_saves_images_to_subdir(self, tmp_path):
        import io

        from PIL import Image

        img = Image.new("RGB", (10, 10), "blue")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        img_bytes = buf.getvalue()

        result = _make_result(
            markdown_text="![img](images/test.png)",
            images={"images/test.png": img_bytes},
        )
        out = tmp_path / "report.md"
        assert ExportService.export(result, out, "markdown")

        img_file = tmp_path / "report_images" / "images" / "test.png"
        assert img_file.exists()
        assert img_file.read_bytes() == img_bytes


class TestExportHtml:
    def test_structured_html_omits_discarded_blocks(self, tmp_path):
        result = _make_result(
            content_list=[
                {"type": "header", "text": "SECRET_HEADER"},
                {"type": "text", "text": "VISIBLE_BODY"},
                {
                    "type": "table",
                    "table": _mixed_table_payload(),
                },
                {"type": "page_number", "text": "SECRET_PAGE_NUMBER"},
                {"type": "footer", "text": "SECRET_FOOTER"},
            ],
            html_text="<p>LOSSY</p>",
        )
        out = tmp_path / "discarded.html"

        assert ExportService.export(result, out, "html")

        content = out.read_text(encoding="utf-8")
        assert "VISIBLE_BODY" in content
        assert "SECRET_HEADER" not in content
        assert "SECRET_PAGE_NUMBER" not in content
        assert "SECRET_FOOTER" not in content

    def test_content_list_tables_replace_lossy_html_tables_in_stable_order(
        self, tmp_path
    ):
        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table": TableModelV1(
                        table_id="first",
                        row_count=1,
                        column_count=1,
                        cells=(
                            TableCellV1(
                                cell_id="first-cell",
                                row=0,
                                column=0,
                                text="CANONICAL_FIRST",
                            ),
                        ),
                    ).to_payload(),
                },
                {"type": "text", "text": "between"},
                {
                    "type": "table",
                    "table_body": (
                        "<table><tr><td>LEGACY_SECOND</td></tr></table>"
                    ),
                },
            ],
            html_text=(
                "<main>"
                "<table><tr><td>LOSSY_FIRST</td></tr></table>"
                "<p>between</p>"
                "<table><tr><td>LOSSY_SECOND</td></tr></table>"
                "</main>"
            ),
        )
        out = tmp_path / "test.html"

        assert ExportService.export(result, out, "html")

        content = out.read_text(encoding="utf-8")
        assert "LOSSY_FIRST" not in content
        assert "LOSSY_SECOND" not in content
        assert content.index("CANONICAL_FIRST") < content.index("LEGACY_SECOND")
        assert "<p>between</p>" in content

    def test_structured_html_parses_each_table_only_once(
        self, tmp_path, monkeypatch
    ):
        import vibeocr.tables.reducer as reducer

        original = reducer.table_model_from_block
        calls = 0

        def tracked(block, *args, **kwargs):
            nonlocal calls
            calls += 1
            return original(block, *args, **kwargs)

        monkeypatch.setattr(reducer, "table_model_from_block", tracked)
        result = _make_result(
            content_list=[
                {"type": "table", "table": _mixed_table_payload()}
            ],
            html_text="<p>LOSSY</p>",
        )

        assert ExportService.export(result, tmp_path / "single-pass.html", "html")
        assert calls == 1

    def test_structured_html_does_not_build_unused_markdown(
        self, tmp_path, monkeypatch
    ):
        import vibeocr.tables.reducer as reducer

        def fail_if_called(*args, **kwargs):
            raise AssertionError("HTML-only export must not build Markdown")

        monkeypatch.setattr(reducer, "table_model_to_markdown", fail_if_called)
        result = _make_result(
            content_list=[
                {"type": "table", "table": _mixed_table_payload()}
            ],
            html_text="<p>LOSSY</p>",
        )
        output = tmp_path / "html-only.html"

        assert ExportService.export(result, output, "html")
        assert "<table" in output.read_text(encoding="utf-8")

    def test_structured_html_matches_shared_projection_for_rich_blocks(
        self, tmp_path
    ):
        from vibeocr.tables.reducer import build_result_projections

        result = _make_result(
            content_list=[
                {
                    "type": "title",
                    "text": "Report",
                    "text_level": 2,
                },
                {
                    "type": "table",
                    "table": _mixed_table_payload(),
                    "table_caption": "Caption",
                    "table_footnote": "Footnote",
                },
                {"type": "list", "list_items": ["One", "Two"]},
                {"type": "code", "code_body": "print(1)"},
            ],
            html_text="<p>LOSSY</p>",
        )
        projections = build_result_projections(result, include_raw=False)
        assert projections is not None
        expected_html = projections[2]
        output = tmp_path / "shared-projection.html"

        assert ExportService.export(result, output, "html")

        assert expected_html in output.read_text(encoding="utf-8")

    def test_exports_html_text(self, tmp_path):
        result = _make_result(html_text="<p>Hello</p>")
        out = tmp_path / "test.html"
        assert ExportService.export(result, out, "html")
        content = out.read_text(encoding="utf-8")
        assert "<!DOCTYPE html>" in content
        assert "<p>Hello</p>" in content

    def test_content_blocks_keep_text_and_table_order_without_html_placeholders(
        self, tmp_path
    ):
        result = _make_result(
            content_list=[
                {"type": "text", "text": "BEFORE"},
                {
                    "type": "table",
                    "table": TableModelV1(
                        table_id="ordered",
                        row_count=1,
                        column_count=1,
                        cells=(
                            TableCellV1(
                                cell_id="ordered-cell",
                                row=0,
                                column=0,
                                text="TABLE",
                            ),
                        ),
                    ).to_payload(),
                },
                {"type": "text", "text": "AFTER"},
            ],
            html_text="<p>LOSSY SUMMARY WITHOUT TABLE PLACEHOLDER</p>",
        )
        out = tmp_path / "ordered.html"

        assert ExportService.export(result, out, "html")

        content = out.read_text(encoding="utf-8")
        assert content.index("BEFORE") < content.index("TABLE") < content.index("AFTER")
        assert "LOSSY SUMMARY" not in content

    def test_structured_image_table_and_list_export_in_order(self, tmp_path):
        result = _make_result(
            content_list=[
                {
                    "type": "image",
                    "img_path": "fig.png",
                    "image_caption": ["FIGURE"],
                },
                {
                    "type": "table",
                    "table": TableModelV1(
                        table_id="image-order",
                        row_count=1,
                        column_count=1,
                        cells=(
                            TableCellV1(
                                cell_id="image-order-cell",
                                row=0,
                                column=0,
                                text="TABLE",
                            ),
                        ),
                    ).to_payload(),
                },
                {"type": "list", "list_items": ["ONE", "TWO"]},
            ],
            html_text="<p>LOSSY</p>",
            images={"fig.png": b"\x89PNG"},
        )
        out = tmp_path / "mixed.html"

        assert ExportService.export(result, out, "html")

        content = out.read_text(encoding="utf-8")
        assert content.index("FIGURE") < content.index("TABLE") < content.index("ONE")
        assert "data:image/png;base64," in content
        assert "<li>ONE</li>" in content

    def test_embeds_base64_images(self, tmp_path):
        import io

        from PIL import Image

        img = Image.new("RGB", (10, 10), "green")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        img_bytes = buf.getvalue()

        result = _make_result(
            html_text='<img src="images/photo.png">',
            images={"images/photo.png": img_bytes},
        )
        out = tmp_path / "test.html"
        assert ExportService.export(result, out, "html")
        content = out.read_text(encoding="utf-8")
        assert "data:image/png;base64," in content

    def test_jpg_mime_type(self, tmp_path):
        result = _make_result(
            html_text='<img src="pic.jpg">',
            images={"pic.jpg": b"\xff\xd8\xff\xe0"},
        )
        out = tmp_path / "test.html"
        assert ExportService.export(result, out, "html")
        content = out.read_text(encoding="utf-8")
        assert "data:image/jpeg;base64," in content


class TestExportDocxExtra:
    def test_legacy_table_uses_native_horizontal_and_vertical_merges(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": (
                        "<table>"
                        '<tr><td rowspan="2">A</td><td colspan="2">B</td></tr>'
                        "<tr><td>C</td><td>D</td></tr>"
                        "</table>"
                    ),
                },
            ],
        )
        out = tmp_path / "test.docx"

        assert ExportService.export(result, out, "docx")

        table = Document(out).tables[0]
        assert table.cell(0, 0)._tc is table.cell(1, 0)._tc
        assert table.cell(0, 1)._tc is table.cell(0, 2)._tc
        assert table.cell(1, 1).text == "C"
        assert table.cell(1, 2).text == "D"

    def test_canonical_table_uses_native_horizontal_and_vertical_merges(
        self, tmp_path
    ):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table": _mixed_table_payload(),
                },
            ],
        )
        out = tmp_path / "test.docx"

        assert ExportService.export(result, out, "docx")

        table = Document(out).tables[0]
        assert table.cell(0, 0)._tc is table.cell(1, 0)._tc
        assert table.cell(0, 1)._tc is table.cell(0, 2)._tc
        assert table.cell(1, 1).text == "C"
        assert table.cell(1, 2).text == "D"

    def test_title_block(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[{"type": "title", "text": "My Title", "level": 2}],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        headings = [
            p
            for p in doc.paragraphs
            if (s := p.style) is not None and (s.name or "").startswith("Heading")
        ]
        assert any("My Title" in h.text for h in headings)

    def test_equation_block(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[{"type": "equation", "text": "a^2+b^2=c^2"}],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        assert any("a^2+b^2=c^2" in p.text for p in doc.paragraphs)

    def test_code_block(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[{"type": "code", "code_body": "print('hi')"}],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        assert any("print" in p.text for p in doc.paragraphs)

    def test_fallback_to_raw_text(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(raw_text="line1\nline2")
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert "line1" in texts
        assert "line2" in texts

    def test_table_caption_and_footnote(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": "<tr><td>X</td></tr>",
                    "table_caption": ["Table 1"],
                    "table_footnote": ["note"],
                },
            ],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Table 1" in all_text
        assert "note" in all_text


class TestExportXlsxExtra:
    def test_legacy_table_preserves_logical_coordinates_and_merges(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": (
                        "<table>"
                        '<tr><td rowspan="2">A</td><td colspan="2">B</td></tr>'
                        "<tr><td>C</td><td>D</td></tr>"
                        "</table>"
                    ),
                },
            ],
        )
        out = tmp_path / "test.xlsx"

        assert ExportService.export(result, out, "xlsx")

        workbook = load_workbook(out)
        sheet = workbook["表格 1"]
        assert sheet["A1"].value == "A"
        assert sheet["B1"].value == "B"
        assert sheet["B2"].value == "C"
        assert sheet["C2"].value == "D"
        assert {str(cell_range) for cell_range in sheet.merged_cells.ranges} == {
            "A1:A2",
            "B1:C1",
        }

    def test_canonical_table_preserves_logical_coordinates_and_merges(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table": _mixed_table_payload(),
                },
            ],
        )
        out = tmp_path / "test.xlsx"

        assert ExportService.export(result, out, "xlsx")

        workbook = load_workbook(out)
        sheet = workbook["表格 1"]
        assert [
            sheet["A1"].value,
            sheet["B1"].value,
            sheet["B2"].value,
            sheet["C2"].value,
        ] == ["A", "B", "C", "D"]
        assert {str(cell_range) for cell_range in sheet.merged_cells.ranges} == {
            "A1:A2",
            "B1:C1",
        }

    def test_title_block(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[{"type": "title", "text": "Report"}],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws = wb.active
        assert ws is not None and ws.title == "文本汇总"

    def test_code_block(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[{"type": "code", "code_body": "x=1"}],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws2 = wb.active
        assert ws2 is not None
        values = [c.value for row in ws2.iter_rows() for c in row if c.value]
        assert any("x=1" in str(v) for v in values)

    def test_fallback_to_raw_text(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(raw_text="hello\nworld")
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws3 = wb.active
        assert ws3 is not None
        values = [c.value for row in ws3.iter_rows() for c in row if c.value]
        assert "hello" in values
        assert "world" in values

    def test_table_only_removes_sheet(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {"type": "table", "table_body": "<tr><td>V</td></tr>"},
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        assert "Sheet" not in wb.sheetnames

    def test_table_caption_and_footnote_are_kept_in_summary_order(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": "<table><tr><td>V</td></tr></table>",
                    "table_caption": ["Caption"],
                    "table_footnote": ["Footnote"],
                }
            ]
        )
        out = tmp_path / "metadata.xlsx"

        assert ExportService.export(result, out, "xlsx")

        workbook = load_workbook(out)
        summary = workbook["文本汇总"]
        values = [row[0] for row in summary.iter_rows(values_only=True)]
        assert values == ["[表格标题] Caption", "[表格脚注] Footnote"]
